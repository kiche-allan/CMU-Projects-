"""
GiPH/PRISM Report Generator
==============================
Reads a scanner.py JSON output file and produces a formatted Word document.

Usage:
  python report.py --input <scan.json> [--output <report.docx>]

Author: Allan Kiche, GiPH/PRISM Research, Carnegie Mellon University
"""

import json
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────

DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
DARK_GREY  = RGBColor(0x40, 0x40, 0x40)
RED        = RGBColor(0xC0, 0x39, 0x2B)
GREEN      = RGBColor(0x1D, 0x6B, 0x2E)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ROW_BLUE   = "EBF3FB"
ROW_WHITE  = "FFFFFF"
HDR_BLUE   = "1F3864"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text: str, level: int = 1, color=None):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color
    elif level == 1:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = MED_BLUE
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = DARK_GREY
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_para(doc, text: str, size: int = 10, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    return p


def add_code(doc, text: str, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = color or DARK_GREY
    return p


def add_table(doc, headers: list, rows: list, col_widths: list = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, HDR_BLUE)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell._tc.get_or_add_tcPr()

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = ROW_BLUE if ri % 2 == 0 else ROW_WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_footer(doc, project_name: str):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Allan Kiche  |  GiPH/PRISM Research  |  Carnegie Mellon University  |  {project_name}")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY


# ─────────────────────────────────────────────────────────────────────────────
# REPORT SECTIONS
# ─────────────────────────────────────────────────────────────────────────────

def write_overview(doc, data: dict):
    add_heading(doc, "Overview", 1)
    add_table(doc,
        headers=["Attribute", "Value"],
        rows=[
            ["Project",              data["project"]],
            ["py3 commit",           data["py3_commit"]],
            ["py2 commit",           data["py2_commit"]],
            ["Files scanned",        str(data["total_files_scanned"])],
            ["Binary origins found", str(len(data["origins"]))],
            ["Total pattern hits",   str(data["stats"]["total_hits"])],
            ["Confirmed bugs",       str(data["stats"]["confirmed"])],
            ["Candidates",           str(data["stats"]["candidates"])],
            ["False positives",      str(data["stats"]["false_positives"])],
            ["False positive rate",  f"{data['stats']['false_positive_rate']}%"],
            ["VERDICT",              _verdict_text(data)],
        ],
        col_widths=[2.0, 4.5]
    )


def _verdict_text(data: dict) -> str:
    n = data["stats"]["confirmed"]
    if n == 0:
        return "CLEAN — No Python 2 to 3 migration bugs found"
    return f"{n} CONFIRMED BUG{'S' if n > 1 else ''} FOUND"


def write_stage1(doc, data: dict):
    add_heading(doc, "Stage 1 — Binary Origin Identification", 1)
    add_para(doc, (
        "All binary data entry points were located. These are the ORIGIN nodes in the data flow model. "
        "Each entry point returns bytes in Python 3 but returned str in Python 2."
    ))

    origins = data.get("origins", [])
    if not origins:
        add_para(doc, "No binary origins found.", italic=True)
        return

    # Group by origin type
    by_type = {}
    for o in origins:
        t = o["origin_type"]
        by_type.setdefault(t, []).append(o)

    for origin_type, items in sorted(by_type.items()):
        add_heading(doc, f"{origin_type.replace('_', ' ').title()} ({len(items)} sites)", 2)
        add_table(doc,
            headers=["File", "Line", "Variable", "Code"],
            rows=[[o["file"], str(o["line"]), o["variable"], o["code_snippet"][:60]] for o in items[:20]],
            col_widths=[2.5, 0.5, 1.0, 2.5]
        )


def write_stage2(doc, data: dict):
    add_heading(doc, "Stage 2 — 40-Pattern Scan Results", 1)

    stats = data["stats"]
    add_table(doc,
        headers=["Pattern Group", "Hits"],
        rows=[[g, str(c)] for g, c in stats.get("hit_rate_by_group", {}).items()],
        col_widths=[4.0, 1.5]
    )

    add_heading(doc, "Hit Classification Summary", 2)
    add_table(doc,
        headers=["Classification", "Count", "Meaning"],
        rows=[
            ["CONFIRMED",      str(stats["confirmed"]),       "Pattern new in py3, taint chain traced to binary origin"],
            ["CANDIDATE",      str(stats["candidates"]),      "Pattern in both codebases or chain partially traced"],
            ["FALSE POSITIVE", str(stats["false_positives"]), "Context proves safety — not a migration bug"],
        ],
        col_widths=[1.5, 0.8, 4.2]
    )


def write_confirmed_bugs(doc, data: dict):
    confirmed = data.get("confirmed_bugs", [])
    if not confirmed:
        add_heading(doc, "Confirmed Bugs", 1)
        add_para(doc, "No confirmed bugs found. See candidates for items requiring manual investigation.", italic=True)
        return

    add_heading(doc, f"Confirmed Bugs ({len(confirmed)})", 1)

    for i, bug in enumerate(confirmed, 1):
        add_heading(doc, f"Bug {i} — {bug['pattern_name']}  ({bug['file']}:{bug['line']})", 2)

        add_table(doc,
            headers=["Field", "Value"],
            rows=[
                ["File",          bug["file"]],
                ["Line",          str(bug["line"])],
                ["Pattern group", bug["group_name"]],
                ["Pattern",       bug["pattern_name"]],
                ["Confidence",    bug["confidence"]],
                ["py2 present",   "Yes" if bug["py2_present"] else "No"],
            ],
            col_widths=[1.5, 5.0]
        )

        add_heading(doc, "Code at sink", 3)
        add_code(doc, bug["code_snippet"])

        if bug.get("py2_snippet"):
            add_heading(doc, "py2 equivalent", 3)
            add_code(doc, bug["py2_snippet"], color=DARK_GREY)

        if bug.get("taint_chain"):
            add_heading(doc, "ORIGIN → FLOW → SINK chain", 3)
            for step in bug["taint_chain"]:
                add_code(doc,
                    f"{step.get('level',''):<20}  {step.get('file','')}:{step.get('line','')}",
                    color=MED_BLUE
                )
                add_code(doc, f"  {step.get('description', '')[:80]}")

        if bug.get("notes"):
            add_para(doc, f"Notes: {bug['notes']}", italic=True, color=DARK_GREY)

        add_divider(doc)


def write_candidates(doc, data: dict):
    candidates = data.get("candidates", [])
    if not candidates:
        return

    add_heading(doc, f"Candidates — Manual Investigation Required ({len(candidates)})", 1)
    add_para(doc, (
        "The following hits were classified as CANDIDATE — the pattern exists in both codebases "
        "or the taint chain could not be fully traced automatically. Each requires manual "
        "Stage 3 investigation and a live crash test to confirm or discard."
    ))

    add_table(doc,
        headers=["File", "Line", "Pattern", "Group", "Confidence", "py2 present"],
        rows=[
            [c["file"], str(c["line"]), c["pattern_name"], c["group_name"],
             c["confidence"], "Yes" if c["py2_present"] else "No"]
            for c in candidates[:30]
        ],
        col_widths=[2.0, 0.5, 1.5, 1.5, 0.8, 0.8]
    )


def write_false_positives(doc, data: dict):
    fp = data.get("false_positives", [])
    if not fp:
        return

    add_heading(doc, f"False Positives ({len(fp)})", 1)
    add_para(doc, (
        "These hits were classified as FALSE POSITIVE — no binary origin was found in the same file "
        "and the pattern was already present in the py2_codebase with the same form."
    ))

    # Show top 20 by pattern group
    add_table(doc,
        headers=["File", "Line", "Pattern", "Group"],
        rows=[
            [f["file"], str(f["line"]), f["pattern_name"], f["group_name"]]
            for f in fp[:20]
        ],
        col_widths=[2.5, 0.5, 2.0, 1.5]
    )

    if len(fp) > 20:
        add_para(doc, f"... and {len(fp) - 20} more false positives not shown.", italic=True)


def write_statistics(doc, data: dict):
    add_heading(doc, "Statistics", 1)
    stats = data["stats"]

    add_table(doc,
        headers=["Metric", "Value"],
        rows=[
            ["Project",                      data["project"]],
            ["Python files scanned",         str(stats["total_files"])],
            ["Binary origins identified",    str(stats["total_origins"])],
            ["Total pattern hits",           str(stats["total_hits"])],
            ["Confirmed bugs",               str(stats["confirmed"])],
            ["Candidates",                   str(stats["candidates"])],
            ["False positives",              str(stats["false_positives"])],
            ["False positive rate",          f"{stats['false_positive_rate']}%"],
        ],
        col_widths=[3.0, 3.5]
    )

    if stats.get("origins_by_type"):
        add_heading(doc, "Binary Origins by Type", 2)
        add_table(doc,
            headers=["Origin type", "Count"],
            rows=[[k, str(v)] for k, v in sorted(stats["origins_by_type"].items(), key=lambda x: -x[1])],
            col_widths=[3.0, 1.5]
        )

    if stats.get("confirmed_by_group"):
        add_heading(doc, "Confirmed Bugs by Pattern Group", 2)
        add_table(doc,
            headers=["Pattern group", "Confirmed bugs"],
            rows=[[k, str(v)] for k, v in sorted(stats["confirmed_by_group"].items(), key=lambda x: -x[1])],
            col_widths=[3.0, 1.5]
        )


# ─────────────────────────────────────────────────────────────────────────────
# MAIN REPORT BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(scan_json_path: str, output_path: str):
    with open(scan_json_path, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = section.right_margin  = Inches(0.9)
    section.top_margin   = section.bottom_margin = Inches(0.9)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    add_footer(doc, data["project"])

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"GiPH/PRISM Migration Bug Scan — {data['project']}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Python 2 → Python 3 Migration Bug Detection — Data Flow Methodology")
    run2.font.size = Pt(12)
    run2.font.color.rgb = MED_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Allan Kiche  |  Joannah Nanjekye (Supervisor)  |  Carnegie Mellon University  |  2026")
    run3.font.size = Pt(9)
    run3.font.color.rgb = LIGHT_GREY

    add_divider(doc)

    # Sections
    write_overview(doc, data)
    add_divider(doc)
    write_stage1(doc, data)
    add_divider(doc)
    write_stage2(doc, data)
    add_divider(doc)
    write_confirmed_bugs(doc, data)
    write_candidates(doc, data)
    add_divider(doc)
    write_false_positives(doc, data)
    add_divider(doc)
    write_statistics(doc, data)

    doc.save(output_path)
    print(f"Report written to: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="GiPH/PRISM Report Generator — JSON scan results to Word document"
    )
    parser.add_argument("--input",  required=True,  help="Input JSON file from scanner.py")
    parser.add_argument("--output", default=None,   help="Output .docx file")
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"ERROR: input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    output = args.output
    if not output:
        stem = Path(args.input).stem
        output = f"{stem}_report.docx"

    generate_report(args.input, output)


if __name__ == "__main__":
    main()
